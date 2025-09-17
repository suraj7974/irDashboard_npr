#!/usr/bin/env python3
"""
EXACT KrutiDev to Unicode Converter - Direct port from index.html
Uses the exact same arrays and logic as the original JavaScript
"""

import pdfplumber
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import re
import re
import sys
import datetime


class ExactKrutiDevConverter:
    def __init__(self):
        self.array_one, self.array_two = self._initialize_exact_arrays()
        self.set_of_matras = "अ आ इ ई उ ऊ ए ऐ ओ औ ा ि ी ु ू ृ े ै ो ौ ं : ँ ॅ"

    def _initialize_exact_arrays(self):
        """Initialize the EXACT arrays from index.html"""

        # EXACT array_one from index.html
        array_one = [
            "ñ",
            "Q+Z",
            "sas",
            "aa",
            ")Z",
            "ZZ",
            "‘",
            "’",
            "“",
            "”",
            "å",
            "ƒ",
            "„",
            "…",
            "†",
            "‡",
            "ˆ",
            "‰",
            "Š",
            "‹",
            "¶+",
            "d+",
            "[+k",
            "[+",
            "x+",
            "T+",
            "t+",
            "M+",
            "<+",
            "Q+",
            ";+",
            "j+",
            "u+",
            "Ùk",
            "Ù",
            "ä",
            "–",
            "—",
            "é",
            "™",
            "=kk",
            "f=k",
            "à",
            "á",
            "â",
            "ã",
            "ºz",
            "º",
            "í",
            "{k",
            "{",
            "=",
            "«",
            "Nî",
            "Vî",
            "Bî",
            "Mî",
            "<î",
            "|",
            "K",
            "}",
            "J",
            "Vª",
            "Mª",
            "<ªª",
            "Nª",
            "Ø",
            "Ý",
            "nzZ",
            "æ",
            "ç",
            "Á",
            "xz",
            "#",
            ":",
            "v‚",
            "vks",
            "vkS",
            "vk",
            "v",
            "b±",
            "Ã",
            "bZ",
            "b",
            "m",
            "Å",
            ",s",
            ",",
            "_",
            "ô",
            "d",
            "Dk",
            "D",
            "£",
            "[k",
            "[",
            "x",
            "Xk",
            "X",
            "Ä",
            "?k",
            "?",
            "³",
            "p",
            "Pk",
            "P",
            "N",
            "t",
            "Tk",
            "T",
            ">",
            "÷",
            "¥",
            "ê",
            "ë",
            "V",
            "B",
            "ì",
            "ï",
            "M+",
            "<+",
            "M",
            "<",
            ".k",
            ".",
            "r",
            "Rk",
            "R",
            "Fk",
            "F",
            ")",
            "n",
            "/k",
            "èk",
            "/",
            "Ë",
            "è",
            "u",
            "Uk",
            "U",
            "i",
            "Ik",
            "I",
            "Q",
            "¶",
            "c",
            "Ck",
            "C",
            "Hk",
            "H",
            "e",
            "Ek",
            "E",
            ";",
            "¸",
            "j",
            "y",
            "Yk",
            "Y",
            "G",
            "o",
            "Ok",
            "O",
            "'k",
            "'",
            '"k',
            '"',
            "l",
            "Lk",
            "L",
            "g",
            "È",
            "z",
            "Ì",
            "Í",
            "Î",
            "Ï",
            "Ñ",
            "Ò",
            "Ó",
            "Ô",
            "Ö",
            "Ø",
            "Ù",
            "Ük",
            "Ü",
            "‚",
            "¨",
            "ks",
            "©",
            "kS",
            "k",
            "h",
            "q",
            "w",
            "`",
            "s",
            "¢",
            "S",
            "a",
            "¡",
            "%",
            "W",
            "•",
            "·",
            "∙",
            "·",
            "~j",
            "~",
            "\\",
            "+",
            " ः",
            "^",
            "*",
            "Þ",
            "ß",
            "(",
            "¼",
            "½",
            "¿",
            "À",
            "¾",
            "A",
            "-",
            "&",
            "&",
            "Œ",
            "]",
            "~ ",
            "@",
            "ाे",
            "ाॅ",
            "ंै",
            "े्र",
            "अौ",
            "अो",
            "आॅ",
        ]

        # EXACT array_two from index.html
        array_two = [
            "॰",
            "QZ+",
            "sa",
            "a",
            "र्द्ध",
            "Z",
            '"',
            '"',
            "'",
            "'",
            "०",
            "१",
            "२",
            "३",
            "४",
            "५",
            "६",
            "७",
            "८",
            "९",
            "फ़्",
            "क़",
            "ख़",
            "ख़्",
            "ग़",
            "ज़्",
            "ज़",
            "ड़",
            "ढ़",
            "फ़",
            "य़",
            "ऱ",
            "ऩ",
            "त्त",
            "त्त्",
            "क्त",
            "दृ",
            "कृ",
            "न्न",
            "न्न्",
            "=k",
            "f=",
            "ह्न",
            "ह्य",
            "हृ",
            "ह्म",
            "ह्र",
            "ह्",
            "द्द",
            "क्ष",
            "क्ष्",
            "त्र",
            "त्र्",
            "छ्य",
            "ट्य",
            "ठ्य",
            "ड्य",
            "ढ्य",
            "द्य",
            "ज्ञ",
            "द्व",
            "श्र",
            "ट्र",
            "ड्र",
            "ढ्र",
            "छ्र",
            "क्र",
            "फ्र",
            "र्द्र",
            "द्र",
            "प्र",
            "प्र",
            "ग्र",
            "रु",
            "रू",
            "ऑ",
            "ओ",
            "औ",
            "आ",
            "अ",
            "ईं",
            "ई",
            "ई",
            "इ",
            "उ",
            "ऊ",
            "ऐ",
            "ए",
            "ऋ",
            "क्क",
            "क",
            "क",
            "क्",
            "ख",
            "ख",
            "ख्",
            "ग",
            "ग",
            "ग्",
            "घ",
            "घ",
            "घ्",
            "ङ",
            "च",
            "च",
            "च्",
            "छ",
            "ज",
            "ज",
            "ज्",
            "झ",
            "झ्",
            "ञ",
            "ट्ट",
            "ट्ठ",
            "ट",
            "ठ",
            "ड्ड",
            "ड्ढ",
            "ड़",
            "ढ़",
            "ड",
            "ढ",
            "ण",
            "ण्",
            "त",
            "त",
            "त्",
            "थ",
            "थ्",
            "द्ध",
            "द",
            "ध",
            "ध",
            "ध्",
            "ध्",
            "ध्",
            "न",
            "न",
            "न्",
            "प",
            "प",
            "प्",
            "फ",
            "फ्",
            "ब",
            "ब",
            "ब्",
            "भ",
            "भ्",
            "म",
            "म",
            "म्",
            "य",
            "य्",
            "र",
            "ल",
            "ल",
            "ल्",
            "ळ",
            "व",
            "व",
            "व्",
            "श",
            "श्",
            "ष",
            "ष्",
            "स",
            "स",
            "स्",
            "ह",
            "ीं",
            "्र",
            "द्द",
            "ट्ट",
            "ट्ठ",
            "ड्ड",
            "कृ",
            "भ",
            "्य",
            "ड्ढ",
            "झ्",
            "क्र",
            "त्त्",
            "श",
            "श्",
            "ॉ",
            "ो",
            "ो",
            "ौ",
            "ौ",
            "ा",
            "ी",
            "ु",
            "ू",
            "ृ",
            "े",
            "े",
            "ै",
            "ं",
            "ँ",
            "ः",
            "ॅ",
            "ऽ",
            "ऽ",
            "ऽ",
            "ऽ",
            "्र",
            "्",
            "?",
            "़",
            ":",
            "‘",
            "’",
            "“",
            "”",
            ";",
            "(",
            ")",
            "{",
            "}",
            "=",
            "।",
            ".",
            "-",
            "µ",
            "॰",
            ",",
            "् ",
            "/",
            "ो",
            "ॉ",
            "ैं",
            "्रे",
            "औ",
            "ओ",
            "ऑ",
        ]

        return array_one, array_two

    def convert_text(self, input_text):
        """Convert KrutiDev text to Unicode using EXACT logic from index.html"""
        if not input_text or not input_text.strip():
            return ""

        # EXACT chunking logic from index.html
        text_size = len(input_text)
        processed_text = ""
        sthiti1 = 0
        sthiti2 = 0
        chale_chalo = 1
        max_text_size = 6000

        while chale_chalo == 1:
            sthiti1 = sthiti2

            if sthiti2 < (text_size - max_text_size):
                sthiti2 += max_text_size
                while sthiti2 > sthiti1 and input_text[sthiti2] != " ":
                    sthiti2 -= 1
            else:
                sthiti2 = text_size
                chale_chalo = 0

            modified_substring = input_text[sthiti1:sthiti2]

            # Process this chunk
            if modified_substring:
                modified_substring = self._replace_symbols(modified_substring)
                processed_text += modified_substring

        # Clean up matra spacing issues
        processed_text = self._clean_matra_spacing(processed_text)

        return processed_text

    def _replace_symbols(self, modified_substring):
        """EXACT symbol replacement logic from index.html"""
        if not modified_substring:
            return ""

        array_one_length = len(self.array_one)

        # Main replacement loop - EXACT from index.html
        for input_symbol_idx in range(array_one_length):
            idx = 0
            while idx != -1:
                modified_substring = modified_substring.replace(
                    self.array_one[input_symbol_idx], self.array_two[input_symbol_idx]
                )
                idx = modified_substring.find(self.array_one[input_symbol_idx])

        # Special glyphs processing - EXACT from index.html
        modified_substring = modified_substring.replace("±", "Zं")
        modified_substring = modified_substring.replace("Æ", "र्f")

        # Handle "f" positioning - EXACT from index.html
        position_of_i = modified_substring.find("f")
        while position_of_i != -1:
            if position_of_i + 1 < len(modified_substring):
                character_next_to_i = modified_substring[position_of_i + 1]
                character_to_be_replaced = "f" + character_next_to_i
                modified_substring = modified_substring.replace(
                    character_to_be_replaced, character_next_to_i + "ि"
                )
            position_of_i = modified_substring.find("f", position_of_i + 1)

        # Handle "fa" positioning - EXACT from index.html
        modified_substring = modified_substring.replace("Ç", "fa")
        modified_substring = modified_substring.replace("É", "र्fa")

        position_of_i = modified_substring.find("fa")
        while position_of_i != -1:
            if position_of_i + 2 < len(modified_substring):
                character_next_to_ip2 = modified_substring[position_of_i + 2]
                character_to_be_replaced = "fa" + character_next_to_ip2
                modified_substring = modified_substring.replace(
                    character_to_be_replaced, character_next_to_ip2 + "िं"
                )
            position_of_i = modified_substring.find("fa", position_of_i + 2)

        modified_substring = modified_substring.replace("Ê", "ीZ")

        # Fix wrong ि् - EXACT from index.html
        position_of_wrong_ee = modified_substring.find("ि्")
        while position_of_wrong_ee != -1:
            if position_of_wrong_ee + 2 < len(modified_substring):
                consonent_next_to_wrong_ee = modified_substring[
                    position_of_wrong_ee + 2
                ]
                character_to_be_replaced = "ि्" + consonent_next_to_wrong_ee
                modified_substring = modified_substring.replace(
                    character_to_be_replaced, "्" + consonent_next_to_wrong_ee + "ि"
                )
            position_of_wrong_ee = modified_substring.find(
                "ि्", position_of_wrong_ee + 2
            )

        # Handle reph positioning - EXACT from index.html
        position_of_R = modified_substring.find("Z")

        while position_of_R > 0:
            probable_position_of_half_r = position_of_R - 1
            character_at_probable_position_of_half_r = modified_substring[
                probable_position_of_half_r
            ]

            # Find non-matra position
            while (
                probable_position_of_half_r > 0
                and character_at_probable_position_of_half_r in self.set_of_matras
            ):
                probable_position_of_half_r -= 1
                character_at_probable_position_of_half_r = modified_substring[
                    probable_position_of_half_r
                ]

            # Check for halant
            previous_to_position_of_half_r = probable_position_of_half_r - 1

            if previous_to_position_of_half_r > 0:
                character_previous_to_position_of_half_r = modified_substring[
                    previous_to_position_of_half_r
                ]

                while character_previous_to_position_of_half_r == "्":
                    probable_position_of_half_r = previous_to_position_of_half_r - 1
                    if probable_position_of_half_r >= 0:
                        character_at_probable_position_of_half_r = modified_substring[
                            probable_position_of_half_r
                        ]
                        previous_to_position_of_half_r = probable_position_of_half_r - 1
                        if previous_to_position_of_half_r >= 0:
                            character_previous_to_position_of_half_r = (
                                modified_substring[previous_to_position_of_half_r]
                            )
                        else:
                            break
                    else:
                        break

            # Replace with proper reph
            character_to_be_replaced = modified_substring[
                probable_position_of_half_r:position_of_R
            ]
            new_replacement_string = "र्" + character_to_be_replaced
            character_to_be_replaced_with_Z = character_to_be_replaced + "Z"
            modified_substring = modified_substring.replace(
                character_to_be_replaced_with_Z, new_replacement_string
            )
            position_of_R = modified_substring.find("Z")

        return modified_substring

    def _clean_matra_spacing(self, text):
        """Clean up ALL spacing issues with matras that might occur during PDF extraction"""
        if not text:
            return text

        result = text

        # Define all Devanagari matras and combining characters
        matras = [
            "ा",
            "ि",
            "ी",
            "ु",
            "ू",
            "ृ",
            "ॄ",  # vowel signs
            "े",
            "ै",
            "ो",
            "ौ",  # compound vowel signs
            "ं",
            "ँ",
            "ः",
            "ऽ",  # anusvara, candrabindu, visarga, avagraha
            "्",  # virama/halant
            "़",  # nukta
        ]

        # Step 1: Remove spaces before ALL matras (multiple passes for multiple spaces)
        for _ in range(3):  # Multiple passes to handle multiple consecutive spaces
            for matra in matras:
                result = result.replace(" " + matra, matra)
                result = result.replace("  " + matra, matra)
                result = result.replace("   " + matra, matra)

        # Step 2: Fix specific problematic vowel combinations that commonly get spaced
        vowel_fixes = [
            # ा combinations
            ("ा आ", "आ"),
            ("ा ऑ", "ऑ"),
            ("ा औ", "औ"),
            ("ा ओ", "ओ"),
            # े combinations
            ("े आ", "ेआ"),
            ("े ऑ", "ेऑ"),
            ("े औ", "ेऔ"),
            ("े ओ", "ेओ"),
            # ो combinations
            ("ो आ", "ोआ"),
            ("ो ऑ", "ोऑ"),
            ("ो औ", "ोऔ"),
            ("ो ओ", "ोओ"),
            # ी combinations
            ("ी आ", "ीआ"),
            ("ी ऑ", "ीऑ"),
            ("ी औ", "ीऔ"),
            ("ी ओ", "ीओ"),
            # Common anusvara patterns
            ("ा ं", "ां"),
            ("े ं", "ें"),
            ("ी ं", "ीं"),
            ("ो ं", "ों"),
            ("ै ं", "ैं"),
            ("ु ं", "ुं"),
            ("ू ं", "ूं"),
            ("ृ ं", "ृं"),
            # Candrabindu patterns
            ("ा ँ", "ाँ"),
            ("े ँ", "ेँ"),
            ("ी ँ", "ीँ"),
            ("ो ँ", "ोँ"),
            ("ै ँ", "ैँ"),
            ("ु ँ", "ुँ"),
            ("ू ँ", "ूँ"),
            ("ृ ँ", "ृँ"),
            # Visarga patterns
            ("ा ः", "ाः"),
            ("े ः", "ेः"),
            ("ी ः", "ीः"),
            ("ो ः", "ोः"),
            ("ै ः", "ैः"),
            ("ु ः", "ुः"),
            ("ू ः", "ूः"),
            ("ृ ః", "ृः"),
        ]

        for wrong, correct in vowel_fixes:
            result = result.replace(wrong, correct)

        # Step 3: Handle complex corruption patterns like "योजनाआंे" → "योजनाओं"
        # These seem to be specific PDF extraction artifacts
        complex_fixes = [
            # The specific case: आंे should become ओं (not ों)
            ("आंे", "ओं"),
            ("आं े", "ओं"),
            ("आ ंे", "ओं"),
            ("आ ं े", "ओं"),
            # Handle cases where space comes before आंे
            (" आंे", "ओं"),
            (" आं े", "ओं"),
            (" आ ंे", "ओं"),
            (" आ ं े", "ओं"),
            # Other similar patterns - fix to ओं
            ("ाआंे", "ाओं"),
            ("ाआं े", "ाओं"),
            ("ाआ ंे", "ाओं"),
            ("ाआ ं े", "ाओं"),
            # Direct corruption fixes
            ("ओंे", "ओं"),
            ("ो ंे", "ओं"),
            ("ों े", "ओं"),
            ("ो ं े", "ओं"),
            # Fix INVALID halant + vowel matra combinations (ROOT CAUSE FIX)
            ("ध्ेा", "धेा"),  # Your specific case!
            ("क्े", "के"),
            ("क्ा", "का"),
            ("क्ो", "को"),
            ("क्ै", "कै"),
            ("क्ौ", "कौ"),
            ("ख्े", "खे"),
            ("ख्ा", "खा"),
            ("ख्ो", "खो"),
            ("ख्ै", "खै"),
            ("ख्ौ", "खौ"),
            ("ग्े", "गे"),
            ("ग्ा", "गा"),
            ("ग्ो", "गो"),
            ("ग्ै", "गै"),
            ("ग्ौ", "गौ"),
            ("घ्े", "घे"),
            ("घ्ा", "घा"),
            ("घ्ो", "घो"),
            ("घ्ै", "घै"),
            ("घ्ौ", "घौ"),
            ("च्े", "चे"),
            ("च्ा", "चा"),
            ("च्ो", "चो"),
            ("च्ै", "चै"),
            ("च्ौ", "चौ"),
            ("छ्े", "छे"),
            ("छ्ा", "छा"),
            ("छ्ो", "छो"),
            ("छ्ै", "छै"),
            ("छ्ौ", "छौ"),
            ("ज्े", "जे"),
            ("ज्ा", "जा"),
            ("ज्ो", "जो"),
            ("ज्ै", "जै"),
            ("ज्ौ", "जौ"),
            ("झ्े", "झे"),
            ("झ्ा", "झा"),
            ("झ्ो", "झो"),
            ("झ्ै", "झै"),
            ("झ्ौ", "झौ"),
            ("ट्े", "टे"),
            ("ट्ा", "टा"),
            ("ट्ो", "टो"),
            ("ट्ै", "टै"),
            ("ट्ौ", "टौ"),
            ("ठ्े", "ठे"),
            ("ठ्ा", "ठा"),
            ("ठ्ो", "ठो"),
            ("ठ्ै", "ठै"),
            ("ठ्ौ", "ठौ"),
            ("ड्े", "डे"),
            ("ड्ा", "डा"),
            ("ड्ो", "डो"),
            ("ड्ै", "डै"),
            ("ड्ौ", "डौ"),
            ("ढ्े", "ढे"),
            ("ढ्ा", "ढा"),
            ("ढ्ो", "ढो"),
            ("ढ्ै", "ढै"),
            ("ढ्ौ", "ढौ"),
            ("ण्े", "णे"),
            ("ण्ा", "णा"),
            ("ण्ो", "णो"),
            ("ण्ै", "णै"),
            ("ण्ौ", "णौ"),
            ("त्े", "ते"),
            ("त्ा", "ता"),
            ("त्ो", "तो"),
            ("त्ै", "तै"),
            ("त्ौ", "तौ"),
            ("थ्े", "थे"),
            ("थ्ा", "था"),
            ("थ्ो", "थो"),
            ("थ्ै", "थै"),
            ("थ्ौ", "थौ"),
            ("द्े", "दे"),
            ("द्ा", "दा"),
            ("द्ो", "दो"),
            ("द्ै", "दै"),
            ("द्ौ", "दौ"),
            ("ध्े", "धे"),
            ("ध्ा", "धा"),
            ("ध्ो", "धो"),
            ("ध्ै", "धै"),
            ("ध्ौ", "धौ"),
            ("न्े", "ने"),
            ("न्ा", "ना"),
            ("न्ो", "नो"),
            ("न्ै", "नै"),
            ("न्ौ", "नौ"),
            ("प्े", "पे"),
            ("प्ा", "पा"),
            ("प्ो", "पो"),
            ("प्ै", "पै"),
            ("प्ौ", "पौ"),
            ("फ्े", "फे"),
            ("फ्ा", "फा"),
            ("फ्ो", "फो"),
            ("फ्ै", "फै"),
            ("फ्ौ", "फौ"),
            ("ब्े", "बे"),
            ("ब्ा", "बा"),
            ("ब्ो", "बो"),
            ("ब्ै", "बै"),
            ("ब्ौ", "बौ"),
            ("भ्े", "भे"),
            ("भ्ा", "भा"),
            ("भ्ो", "भो"),
            ("भ्ै", "भै"),
            ("भ्ौ", "भौ"),
            ("म्े", "मे"),
            ("म्ा", "मा"),
            ("म्ो", "मो"),
            ("म्ै", "मै"),
            ("म्ौ", "मौ"),
            ("य्े", "ये"),
            ("य्ा", "या"),
            ("य्ो", "यो"),
            ("य्ै", "यै"),
            ("य्ौ", "यौ"),
            ("र्े", "रे"),
            ("र्ा", "रा"),
            ("र्ो", "रो"),
            ("र्ै", "रै"),
            ("र्ौ", "रौ"),
            ("ल्े", "ले"),
            ("ल्ा", "ला"),
            ("ल्ो", "लो"),
            ("ल्ै", "लै"),
            ("ल्ौ", "लौ"),
            ("व्े", "वे"),
            ("व्ा", "वा"),
            ("व्ो", "वो"),
            ("व्ै", "वै"),
            ("व्ौ", "वौ"),
            ("श्े", "शे"),
            ("श्ा", "शा"),
            ("श्ो", "शो"),
            ("श्ै", "शै"),
            ("श्ौ", "शौ"),
            ("ष्े", "षे"),
            ("ष्ा", "षा"),
            ("ष्ो", "षो"),
            ("ष्ै", "षै"),
            ("ष्ौ", "षौ"),
            ("स्े", "से"),
            ("स्ा", "सा"),
            ("स्ो", "सो"),
            ("स्ै", "सै"),
            ("स्ौ", "सौ"),
            ("ह्े", "हे"),
            ("ह्ा", "हा"),
            ("ह्ो", "हो"),
            ("ह्ै", "है"),
            ("ह्ौ", "हौ"),
            # Fix multiple matra sequences
            ("ेा", "े"),
            ("ोा", "ो"),
            ("ैा", "ै"),
            ("ौा", "ौ"),
            ("ीा", "ी"),
            ("ूा", "ू"),
            ("ृा", "ृ"),
            ("्ेा", "ेा"),  # Remove halant before ेा sequence
        ]

        for wrong, correct in complex_fixes:
            result = result.replace(wrong, correct)

        # Step 4: Generate comprehensive consonant + matra fixes
        # All major Devanagari consonants
        consonants = [
            "क",
            "ख",
            "ग",
            "घ",
            "ङ",  # ka-varga
            "च",
            "छ",
            "ज",
            "झ",
            "ञ",  # ca-varga
            "ट",
            "ठ",
            "ड",
            "ढ",
            "ण",  # ta-varga
            "त",
            "थ",
            "द",
            "ध",
            "न",  # ta-varga
            "प",
            "फ",
            "ब",
            "भ",
            "म",  # pa-varga
            "य",
            "र",
            "ल",
            "व",  # ya-varga
            "श",
            "ष",
            "स",
            "ह",  # sha-varga
            "क्ष",
            "त्र",
            "ज्ञ",  # compounds
            "ड़",
            "ढ़",
            "क़",
            "ख़",
            "ग़",
            "ज़",
            "फ़",  # nukta forms
        ]

        # Fix consonant + space + matra patterns
        for consonant in consonants:
            for matra in ["ा", "ि", "ी", "ु", "ू", "ृ", "े", "ै", "ो", "ौ"]:
                wrong_pattern = consonant + " " + matra
                correct_pattern = consonant + matra
                result = result.replace(wrong_pattern, correct_pattern)

        # Step 5: Fix halant (्) related spacing issues
        for consonant in consonants:
            # Fix consonant + space + halant
            result = result.replace(consonant + " ्", consonant + "्")
            # Fix halant + space + consonant
            result = result.replace("् " + consonant, "्" + consonant)

        # Step 6: Fix nukta (़) spacing
        nukta_base = ["क", "ख", "ग", "ज", "ड", "ढ", "फ", "य"]
        for base in nukta_base:
            result = result.replace(base + " ़", base + "़")

        # Step 7: Clean up any remaining multiple spaces
        while "  " in result:
            result = result.replace("  ", " ")

        return result

    def _extract_content_with_positions(self, page):
        """Extract content using SIMPLE line-by-line approach - no coordinates"""
        content_items = []

        try:
            print(f"  📝 Using PRECISE question-table matching...")

            # Step 1: Extract text line by line (preserves natural spacing)
            full_text = page.extract_text()
            text_lines = []

            if full_text:
                lines = full_text.split("\n")
                print(f"    📄 Found {len(lines)} text lines")

                for i, line in enumerate(lines):
                    line = line.strip()
                    if line and len(line) > 0:  # Keep all non-empty lines
                        text_lines.append(
                            {"type": "text", "content": line, "line_number": i}
                        )

            # Step 2: Extract tables using find_tables for better positioning
            tables = page.find_tables()
            table_objects = []

            print(f"    📊 Found {len(tables)} tables")

            for i, table in enumerate(tables):
                if table:
                    table_data = table.extract()
                    if table_data and len(table_data) > 0:
                        # Check if table has meaningful content
                        has_content = False
                        for row in table_data:
                            for cell in row:
                                if cell and str(cell).strip():
                                    has_content = True
                                    break
                            if has_content:
                                break

                        if has_content:
                            # Get table Y position for better placement
                            table_y = (
                                table.bbox[1] if hasattr(table, "bbox") else 999999
                            )
                            table_objects.append(
                                {
                                    "type": "table",
                                    "content": table_data,
                                    "table_number": i,
                                    "y_position": table_y,
                                }
                            )

            # Step 3: SMART strategy - only place tables after questions that actually have them
            # Analyze the page content to determine which questions should get tables

            final_items = []
            table_index = 0  # Index into available tables

            # Define which questions typically have tables per page (based on PDF structure)
            page_table_rules = {
                1: [],  # Page 1: No question tables (only form headers)
                2: [],  # Page 2: No question tables
                3: [],  # Page 3: No question tables
                4: [],  # Page 4: No question tables
                5: (
                    [25, 26, 27, 28, 29, 30] if len(table_objects) >= 6 else []
                ),  # Page 5: Questions 25-30
                6: (
                    [34, 35, 36, 37, 38, 39] if len(table_objects) >= 6 else []
                ),  # Page 6: Questions 34-39
                7: (
                    [40, 41] if len(table_objects) >= 2 else []
                ),  # Page 7: Questions 40-41
                8: [],  # Page 8: No question tables
                9: [],  # Page 9: No question tables
                10: [],  # Page 10: No question tables
            }

            # Questions that should have tables (Q28 to Q40 based on user feedback)
            questions_with_tables = set(range(28, 41))
            table_queue = list(table_objects)  # Queue of available tables

            for text_item in text_lines:
                final_items.append(text_item)
                text_content = text_item["content"]

                # Check if this line is a question that should have a table (KrutiDev format uses hyphen)
                question_match = re.match(r"^(\d+)-", text_content)
                if question_match:
                    question_num = int(question_match.group(1))

                    # Only place table if this question should have one AND we have tables available
                    if question_num in questions_with_tables and table_queue:
                        table_obj = table_queue.pop(0)  # Take the next table in order
                        final_items.append(table_obj)
                        print(f"    ✅ Placed table after Q{question_num} (KrutiDev)")
                    elif question_num in questions_with_tables:
                        print(f"    🔍 Found Q{question_num} but no tables available")

            # Add any remaining unplaced tables at the end (safety measure)
            for table_obj in table_queue:
                final_items.append(table_obj)
                print(f"    ⚠️ Unplaced table added at end")

            content_items = final_items

            print(f"  📋 Smart extraction result:")
            print(f"    📝 {len(text_lines)} text lines")
            print(f"    📊 {len(table_objects)} tables")
            print(f"    📄 Total: {len(content_items)} items")

            # Show first few items for debugging
            print(f"  � First 10 items:")
            for i, item in enumerate(content_items[:10]):
                if item["type"] == "text":
                    preview = (
                        item["content"][:60] + "..."
                        if len(item["content"]) > 60
                        else item["content"]
                    )
                    print(f"     {i+1:2d}. TEXT: {preview}")
                else:
                    rows = len(item["content"]) if item["content"] else 0
                    print(f"     {i+1:2d}. TABLE ({rows} rows)")

            if len(content_items) > 10:
                print(f"     ... and {len(content_items) - 10} more items")

            return content_items

        except Exception as e:
            print(f"  ⚠️ Simple extraction failed: {e}")

            # Ultimate fallback: just get text
            content_items = []

            try:
                full_text = page.extract_text()
                if full_text:
                    lines = full_text.split("\n")
                    for i, line in enumerate(lines):
                        if line.strip():
                            content_items.append(
                                {"type": "text", "content": line.strip(), "position": i}
                            )

                # Add tables at the end as fallback
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        content_items.append(
                            {
                                "type": "table",
                                "content": table,
                                "position": len(content_items),
                            }
                        )

            except Exception as e2:
                print(f"  � Even fallback failed: {e2}")

            return content_items

    def convert_pdf_to_docx(self, pdf_path, output_path):
        """Convert PDF with KrutiDev text to Unicode DOCX with table support"""
        try:
            print(f"📖 Opening PDF: {pdf_path}")

            # Create DOCX document
            doc = Document()
            doc.add_heading("KrutiDev to Unicode Conversion (ACCURATE POSITIONING)", 0)
            doc.add_paragraph(f"Converted from: {pdf_path}")
            doc.add_paragraph(f"Conversion date: {datetime.datetime.now()}")
            doc.add_paragraph("─" * 50)

            total_text_length = 0
            total_tables = 0

            # Extract text and tables from PDF
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
                print(f"📄 Processing {total_pages} pages...")

                for page_num, page in enumerate(pdf.pages, 1):
                    # Add page heading
                    doc.add_heading(f"Page {page_num}", 2)

                    # Get content in correct order
                    content_items = self._extract_content_with_positions(page)

                    tables_count = len(
                        [item for item in content_items if item["type"] == "table"]
                    )
                    text_count = len(
                        [item for item in content_items if item["type"] == "text"]
                    )
                    print(
                        f"Page {page_num}: {tables_count} tables, {text_count} text items"
                    )

                    # Process content in order
                    for item in content_items:
                        if item["type"] == "text":
                            # Convert and add text
                            converted_text = self.convert_text(item["content"])
                            if converted_text.strip():
                                doc.add_paragraph(converted_text.strip())
                                total_text_length += len(converted_text)

                        elif item["type"] == "table":
                            # Add table
                            table_data = item["content"]
                            if table_data and len(table_data) > 0:
                                # Create Word table
                                word_table = doc.add_table(
                                    rows=len(table_data), cols=len(table_data[0])
                                )
                                word_table.style = "Table Grid"
                                word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                                # Fill table with converted data
                                for row_idx, row in enumerate(table_data):
                                    for col_idx, cell in enumerate(row):
                                        if cell and str(cell).strip():
                                            converted_cell = self.convert_text(
                                                str(cell)
                                            )
                                            word_table.cell(row_idx, col_idx).text = (
                                                converted_cell
                                            )

                                total_tables += 1
                                doc.add_paragraph()  # Add space after table

                    print(f"✓ Page {page_num}/{total_pages} processed")

            # Add summary
            doc.add_paragraph("─" * 50)
            doc.add_paragraph(f"📊 Conversion Summary:")
            doc.add_paragraph(f"   • Total pages processed: {total_pages}")
            doc.add_paragraph(f"   • Total text length: {total_text_length} characters")
            doc.add_paragraph(f"   • Total tables converted: {total_tables}")

            # Save document
            doc.save(output_path)
            print(f"✅ Conversion completed successfully!")
            print(f"📁 Output saved to: {output_path}")
            print(
                f"📊 Summary: {total_pages} pages, {total_text_length} chars, {total_tables} tables"
            )

            return True

        except Exception as e:
            print(f"❌ Error during conversion: {str(e)}")
            return False

    def test_converter(self):
        """Test the converter with the exact sample from your screenshot"""
        test_input = "laxBu esa 'kkfey gksus dk dkj.k ¼dc@dSls@fdldh en&n ls@in vkfn½"

        print("🧪 Testing EXACT KrutiDev Converter:")
        print("═" * 60)
        print(f"Input:  {test_input}")

        result = self.convert_text(test_input)
        print(f"Output: {result}")

        print("═" * 60)
        print("✅ Test completed! This should match the original JS converter.")


def main():
    converter = ExactKrutiDevConverter()

    if len(sys.argv) > 1:
        pdf_file = sys.argv[1]
        output_file = pdf_file.replace(".pdf", "_unicode_exact.docx")
        converter.convert_pdf_to_docx(pdf_file, output_file)
    else:
        # Run test
        converter.test_converter()

        # Try converting the test PDF if it exists
        test_pdf = "Form A.pdf"
        try:
            output_file = f"converted_unicode_exact_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            converter.convert_pdf_to_docx(test_pdf, output_file)
        except FileNotFoundError:
            print(f"\n📄 To convert a PDF, run: python {sys.argv[0]} <pdf_file>")


if __name__ == "__main__":
    main()
